#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Update the product .docx specs with the HELIX Agent Cast roster (the team + the
new roles). Appends a "הצוות (Agent Cast)" section to each product's existing
.docx — non-destructive (keeps all current content) and idempotent (skips a file
that already has the section). Source of truth: PRODUCTS/HELIX-AGENT-CAST.md.

Run:  python scripts/update-docx.py
Needs: python-docx  (pip install python-docx)
"""
import os
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
PRODUCTS = os.path.join(HERE, "..", "PRODUCTS")
MARKER = "הצוות (Agent Cast)"

# filename (in PRODUCTS/) -> list of (name, role, one-liner)
ROSTERS = {
    "04-seo-agent.docx": [  # Rank
        ("עמנואל", "חוקר מילים", "מוצא מילים וישויות, ובודק חפיפה לדפים קיימים."),
        ("דן", "הכתב", "כותב את המאמר."),
        ("אלון", "העורך שפוסל", "פוסל טענה לא-מדויקת, קניבליזציה ותוכן ש-AI לא יצטט."),
        ("מאיה", "העורכת", "משכתבת לפי אלון."),
        ("CHIEF", "המנהל", "מתאם בין הצוות."),
    ],
    "07-growth-doctor.docx": [
        ("ליה", "האנליסטית", "בודקת אם הנשירה מובהקת או רעש."),
        ("דן", "כותב התיקון", "מנסח את ההמלצה הקונקרטית."),
        ("אלון", "מבקר הביצוע", "עוצר תיקון על אבחון חלש לפני שהוא רץ."),
        ("מאיה", "העורכת", "מחדדת את ההמלצה."),
        ("CHIEF", "המנהל", "מתאם בין הצוות."),
    ],
    "01-marketing-ops-hub.docx": [
        ("עמנואל", "חוקר ההקשר", "קורא את הפוסט ומסמן אם הוא רגיש."),
        ("דן", "כותב התגובה", "מנסח בקול המותג."),
        ("אלון", "מבקר brand-safety", "פוסל תגובה ספאמית, DM לא-הולם והזזת-תקציב מסוכנת."),
        ("מאיה", "העורכת", "משכתבת לפי אלון."),
        ("CHIEF", "המנהל", "מתאם בין הצוות."),
    ],
    "03-helix-sdr-bdr-bot.docx": [
        ("עמנואל", "ההעשרה", "אוסף מידע על הליד."),
        ("ליה", "מאמתת המקור", "מצליבה כל נתון מול המקור."),
        ("דן", "הכתב", "מנסח את הפנייה."),
        ("מאיה", "העורכת", "משכתבת עד שזה אנושי."),
        ("אלון", "מבקר השליחה", "בודק כל פנייה לפני שליחה, כולל שלא תיפול לספאם."),
        ("נעמי", "הרכזת", "מחליטה למי לפנות ואיזה פולואפ."),
        ("CHIEF", "המנהל", "מתאם בין הצוות."),
    ],
    "12-helix-shop.docx": [
        ("עמנואל", "מומחה המוצר", "שולף מהקטלוג עובדות בלבד."),
        ("דן", "המוכר", "עונה, ממליץ ומוסיף לעגלה."),
        ("אלון", "בודק העובדות", "מוודא שכל מחיר ומלאי קיימים לפני שליחה."),
        ("CHIEF", "המנהל", "מתאם בין הצוות."),
    ],
    "10-helix-meeting.docx": [
        ("עמנואל", "מחלץ ההתחייבויות", "מוציא מהתמליל הבטחות עם ציטוט."),
        ("דן", "כותב המייל", "מנסח מייל-סיכום מההתחייבויות בלבד."),
        ("אלון", "המאמת", "מוודא שכל דבר במייל באמת נאמר."),
        ("מאיה", "העורכת", "משכתבת אם צריך."),
        ("נעמי", "הרכזת", "מייחסת משימות וקובעת את הפגישה הבאה."),
        ("CHIEF", "המנהל", "מתאם בין הצוות."),
    ],
    "02-HELIX-dashboards.docx": [
        ("ליה", "האנליסטית", "אוספת KPIs ומדדים שמידרדרים."),
        ("דן", "המספר", "כותב את הסיכום."),
        ("אלון", "המבקר", "פוסל מספר שלא קיים בדאטה."),
        ("מאיה", "העורכת", "משכתבת מהעובדות בלבד."),
        ("CHIEF", "המנהל", "מתאם בין הצוות."),
    ],
    "05-reputation-agent.docx": [
        ("עמנואל", "המנטר", "סורק ביקורות, סושיאל ומנועי-AI."),
        ("דן", "המגיב", "מנסח תגובות ותוכן חיובי."),
        ("אלון", "השער האתי", "מסרב לפעול נגד עיתונות לגיטימית או לזייף ביקורות."),
        ("CHIEF", "המנהל", "מתאם בין הצוות."),
    ],
    "08-helix-sign-forms.docx": [
        ("עמנואל", "מנתח הבקשה", "מפרק את הבקשה לצדדים ותנאים."),
        ("דן", "המנסח", "כותב את החוזה."),
        ("אלון", "מבקר הסעיפים", "מסמן סעיף חד-צדדי או חסר."),
        ("מאיה", "מתקן החוזה", "מוסיף סעיפים חסרים."),
        ("CHIEF", "המנהל", "מתאם בין הצוות."),
    ],
}


def has_marker(doc: Document) -> bool:
    return any(MARKER in p.text for p in doc.paragraphs)


def main() -> None:
    updated, skipped, missing = 0, 0, 0
    for fname, roster in ROSTERS.items():
        path = os.path.join(PRODUCTS, fname)
        if not os.path.exists(path):
            print(f"  missing: {fname}")
            missing += 1
            continue
        doc = Document(path)
        if has_marker(doc):
            print(f"  skip (already has cast): {fname}")
            skipped += 1
            continue
        doc.add_heading(MARKER, level=1)
        doc.add_paragraph("הצוות שעובד בשבילך — מחלקה של סוכנים, כל אחד עם תפקיד, עם מבקר שבודק את העבודה לפני שהיא יוצאת:")
        for name, role, line in roster:
            p = doc.add_paragraph()
            r = p.add_run(f"•  {name} — {role}: ")
            r.bold = True
            p.add_run(line)
        doc.save(path)
        print(f"  updated: {fname}")
        updated += 1
    print(f"\nDone. updated={updated} skipped={skipped} missing={missing}")


if __name__ == "__main__":
    main()
